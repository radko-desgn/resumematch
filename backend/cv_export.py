"""Export a tailored CV (ATS-safe markdown) to PDF or DOCX.

The prompt constrains the CV to a small markdown subset — `#`, `##`, `-`, and
`**bold**` — so a focused parser is more predictable here than a full markdown
engine, and keeps the output ATS-friendly (no tables, columns, or graphics).
"""

from __future__ import annotations

import html
import io
import re

_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _inline(text: str) -> str:
    return _BOLD.sub(r"<strong>\1</strong>", html.escape(text))


def markdown_to_html(md: str) -> str:
    """Convert the constrained CV markdown subset into semantic HTML."""
    out: list[str] = []
    in_list = False

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            close_list()
            continue
        if line.startswith("## "):
            close_list()
            out.append(f"<h2>{_inline(line[3:])}</h2>")
        elif line.startswith("# "):
            close_list()
            out.append(f"<h1>{_inline(line[2:])}</h1>")
        elif line.lstrip().startswith(("- ", "* ")):
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{_inline(line.lstrip()[2:])}</li>")
        else:
            close_list()
            out.append(f"<p>{_inline(line)}</p>")
    close_list()
    return "\n".join(out)


def _cv_document(md: str) -> str:
    """Wrap the CV in a clean, print-ready, ATS-safe page."""
    return f"""<!doctype html><html><head><meta charset="utf-8">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap" rel="stylesheet">
<style>
  @page {{ margin: 16mm 15mm; }}
  body {{ font-family: Inter, system-ui, sans-serif; color:#0A0A0A; font-size:10.5pt; line-height:1.5; margin:0; }}
  h1 {{ font-size:20pt; margin:0 0 2px; letter-spacing:-0.02em; }}
  h2 {{ font-size:11pt; margin:16px 0 6px; text-transform:uppercase; letter-spacing:.08em;
        border-bottom:1px solid #DDD; padding-bottom:3px; }}
  p  {{ margin:4px 0; }}
  ul {{ margin:4px 0 8px; padding-left:16px; }}
  li {{ margin:3px 0; }}
  strong {{ font-weight:600; }}
</style></head><body>
{markdown_to_html(md)}
</body></html>"""


def cv_to_pdf(md: str) -> bytes:
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(_cv_document(md), wait_until="networkidle")
        pdf = page.pdf(format="A4", print_background=True)
        browser.close()
    return pdf


def cv_to_docx(md: str) -> bytes:
    """Build a real .docx — headings and bullets as native Word styles."""
    import docx

    doc = docx.Document()

    def add_runs(paragraph, text: str) -> None:
        # split on **bold** so Word gets real bold runs, not literal asterisks
        for part in re.split(r"(\*\*.+?\*\*)", text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            else:
                paragraph.add_run(part)

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            add_runs(doc.add_paragraph(style="List Bullet"), line.lstrip()[2:])
        else:
            add_runs(doc.add_paragraph(), line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
